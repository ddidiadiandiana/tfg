import json
import re
from datetime import date
from typing import Any, Dict, List, Tuple

import streamlit as st
from collections import defaultdict
from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject, TextStringObject, BooleanObject, DictionaryObject

from langchain.chains import LLMChain

def get_prompt_template(extension: str) -> Tuple[str, str]:
    """
    Returns system and user prompt templates based on the file extension.

    :param extension: File extension indicating document type.
    :return: Tuple containing (system_template, user_template) strings for prompting the language model.
    """    
    if extension == ".docx":
        system_template = """
You are an expert assistant in filling out forms accurately and professionally.
""".strip()

        user_template = """
Today's date is {today}.

Instructions:
- Use the user input to complete all required fields.
- Match expected formats carefully (name order eg "Surname, Name"; case sensitivity and accents eg "gran via" → "Gran Vía"; date formats eg "1 de enero de 2025" → "01/01/2025").
- Use uppercase letters in IDs, bank accounts, and similar fields.
- Do not repeat currency symbols or labels if they already appear in the document.
- Use {today} date if the user mentions "today"/"date" (case-insensitive and in any language). Infer the end date if the user provides a start date and a duration.
- Ask for any missing info using natural language.

Given the following fields:
{fields}

And the user input:
<<< {user_input} >>>

Return only a JSON dictionary with keys that exactly match the field names.
""".strip()

    elif extension == ".pdf":
        system_template = """
You are an expert assistant in filling out forms accurately and professionally.
""".strip()
        
        user_template = """
Today's date is {today}.

Instructions:
- Use the user input to complete all required fields.
- Each field has a name, dimensions, and position.
- Match expected formats carefully (name order eg "Surname, Name"; case sensitivity and accents eg "gran via" → "Gran Vía"; date formats eg "1 de enero de 2025" → "01/01/2025").
- Use uppercase letters in IDs, bank accounts, and similar fields.
- Do not repeat currency symbols or labels if they already appear in the document.
- Use {today} date if the user mentions "today"/"date" (case-insensitive and in any language). Infer the end date if the user provides a start date and a duration.
- Ask for any missing info using natural language.

Fields: {fields}

User input: <<< {user_input} >>>

Return only a JSON dictionary with keys that exactly match the field names.
""".strip()

    return system_template, user_template

def extract_fields(file: str, extension: str) -> List[Dict[str, Any]]:
    """
    Extracts candidate form fields from a DOCX or PDF file.

    :param file: Path to the input file.
    :param extension: File extension indicating document type.
    :return: List of dictionaries representing detected fields and their metadata.
    """
    if extension == ".docx":
        doc = Document(file)
        fields = []
    
        # Paragraph fields
        for paragraph_i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
        
            if "___" in text:
                fragments = re.split(r"_{3,}", text)
            else:
                fragments = [text]
        
            for frag in fragments:
                frag = frag.strip()
                if not frag:
                    continue
                    
                if "\n" in frag:
                    frag = frag.split("\n")[-1].strip()
                if ":" in frag and len(frag.split(":")) > 1:
                    frag = frag.split(":")[0].strip() + ":"
        
                word_count = len(frag.split())
                if ":" in frag or "-" in frag or word_count <= 10:
                    fields.append({
                        "name": frag,
                        "location": "paragraph",
                        "paragraph": paragraph_i,
                        "append_style": "inline"
                    })
    
        # Table fields
        for table_i, table in enumerate(doc.tables):
            for row_i, row in enumerate(table.rows):
                for column_i, cell in enumerate(row.cells):
                    field_name = cell.text.strip()
                    if field_name and (":" in field_name or "-" in field_name or len(field_name.split()) <= 10):
                        if column_i + 1 < len(row.cells) and not row.cells[column_i + 1].text.strip():
                            append_style = "cell_right"
                        elif row_i + 1 < len(table.rows) and not table.rows[row_i + 1].cells[column_i].text.strip():
                            append_style = "cell_below"
                        else:
                            append_style = "inline"
    
                        fields.append({
                            "name": field_name,
                            "location": "table",
                            "table": table_i,
                            "row": row_i,
                            "column": column_i,
                            "append_style": append_style
                        })

        grouped = defaultdict(list)
        for f in fields:
            if f["location"] == "table":
                grouped[f["name"]].append(f)

        cleaned_fields = []
        for f in fields:
            if f["location"] != "table":
                cleaned_fields.append(f)
                continue

            items = grouped[f["name"]]
            if len(items) > 1:
                non_inline = [x for x in items if x["append_style"] != "inline"]  # cell_right or cell_below
                if f["append_style"] == "inline" and non_inline:
                    continue
            cleaned_fields.append(f)

        fields = cleaned_fields
        
    elif extension == ".pdf":
        reader = PdfReader(file)
        fields = []
    
        for page_num, page in enumerate(reader.pages, start=1):
            annots = page.get("/Annots")
            if annots is not None:
                for annotation in list(annots):
                    obj = annotation.get_object()
                    name = str(obj.get("/T")) if isinstance(obj.get("/T"), TextStringObject) else ""
                    value = str(obj.get("/V")) if isinstance(obj.get("/V"), TextStringObject) else ""
                    rect = obj.get("/Rect", [0, 0, 0, 0])
    
                    try:
                        width = float(rect[2]) - float(rect[0])
                        height = float(rect[3]) - float(rect[1])
                    except Exception:
                        width, height = 0, 0
    
                    fields.append({
                        "name": name,
                        "default_value": value,
                        "width": width,
                        "height": height,
                        "position": rect,
                        "page": page_num
                    })

    return fields

def extract_json_dict(text: str) -> str:
    """
    Extracts the first JSON-like dictionary from a string.

    :param text: The input string potentially containing a JSON object.
    :return: A string containing a JSON block, or an empty string if none is found.
    """

    stack = []
    start = None

    for i, char in enumerate(text):
        if char == '{':
            if not stack:
                start = i
            stack.append('{')
        elif char == '}':
            if stack:
                stack.pop()
                if not stack and start is not None:
                    return text[start:i+1]

    return ""

def infer_data(fields: List[Dict[str, Any]], user_input: str, chain: LLMChain) -> Dict[str, str]:
    """
    Infers field values from user input using a language model chain.
    
    :param fields: List of dictionaries containing field names and their properties.
    :param user_input: Text input providing relevant personal information from the user.
    :param chain: LangChain LLMChain used to invoke the language model.
    :return: Dictionary mapping field names to inferred values.
    """
    fields = [field["name"] for field in fields]
    
    prompt_vars = {
        "today": date.today().strftime("%d/%m/%Y"),
        "fields": json.dumps(fields, indent=2, ensure_ascii=False),
        "user_input": user_input
    }

    response = chain.invoke(prompt_vars).content.strip()

    tokens = ["na", "n/a", "none", "null"]
    pattern = r'\b(?:' + '|'.join(map(re.escape, tokens)) + r')\b'
    response = re.sub(pattern, '', response, flags=re.IGNORECASE)

    try:
        return json.loads(response)
    except json.JSONDecodeError:
        extracted_json = extract_json_dict(response)
        try:
            return json.loads(extracted_json)
        except json.JSONDecodeError:
            st.error("No se ha podido interpretar la respuesta del modelo como un JSON válido. El proceso se ha detenido.", icon=":material/error:")
            st.stop()

def fill_form(inputf: str, outputf: str, data: Dict[str, str], fields: List[Dict[str, Any]], extension: str) -> None:
    """
    Fills document form fields with provided data.
    
    :param inputf: Path to the input file to be filled.
    :param outputf: Path to save the filled output document.
    :param data: Dictionary associating field names with fill values.
    :param fields: List of field metadata extracted from the document.
    :param extension: File extension indicating document type.
    :param mapping: Dictionary mapping original field names to inferred field names.
    """
    if extension == ".docx":
        doc = Document(inputf)

        for field in fields:
            extracted_key = field["name"]

            if extracted_key not in data:
                continue

            value = data.get(extracted_key, "")
            if not value:
                continue

            # Fill paragraphs
            if field["location"] == "paragraph":
                paragraph = doc.paragraphs[field["paragraph"]]
                if extracted_key in paragraph.text:
                    if "___" in paragraph.text:  # tiene guiones bajos después
                        pattern = re.escape(extracted_key) + r"[ _]*_{3,}"
                        paragraph.text = re.sub(pattern, f"{extracted_key} {value} ", paragraph.text)
                    else:
                        paragraph.text = paragraph.text.replace(extracted_key, f"{extracted_key} {value}")
    
            # Fill tables
            elif field["location"] == "table":
                table = doc.tables[field["table"]]
                row, column = field["row"], field["column"]

                if field["append_style"] == "cell_right":
                    table.cell(row, column + 1).text = value
                elif field["append_style"] == "cell_below":
                    table.cell(row + 1, column).text = value
                else:
                    table.cell(row, column).text = f"{extracted_key} {value}"
    
        doc.save(outputf)
        
    elif extension == ".pdf":
        reader = PdfReader(inputf)

        for page in reader.pages:
            annots = page.get("/Annots")
            if annots is not None:
                for annotation in list(annots):
                    obj = annotation.get_object()
                    name = obj.get("/T")
    
                    if isinstance(name, TextStringObject):
                        field_name = str(name)
                    else:
                        field_name = ""
    
                    value = data.get(field_name, "")
    
                    if value:
                        obj.update({NameObject("/V"): TextStringObject(str(value))})
    
        # if "/AcroForm" in reader.trailer["/Root"]:
        #     acro_form = reader.trailer["/Root"]["/AcroForm"]
        #     acro_form.update({NameObject("/NeedAppearances"): BooleanObject(True)})
        # else:
        #     reader.trailer["/Root"][NameObject("/AcroForm")] = DictionaryObject({
        #         NameObject("/NeedAppearances"): BooleanObject(True)
        #     })
        
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        # writer._root_object.update({
        #     NameObject("/AcroForm"): reader.trailer["/Root"]["/AcroForm"]
        # })
    
        writer.write(outputf)